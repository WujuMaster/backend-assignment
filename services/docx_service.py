import json
from os import path
from typing import Union

import docx2txt
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell

from utils.helper import prepare_output_dir


class DocxExtractor:
    def __init__(self, data_dir: str, output_dir: str):
        self.data_dir = data_dir
        self.output_dir = output_dir

    def extract_images(self, filename: str, img_dir: str):
        """Extract images from the docx file
        Args:
            filename (str): Name of the docx file
            img_dir (str): Directory to save the images
        """
        docx2txt.process(path.join(self.data_dir, filename), img_dir)

    def extract_formatting_and_convert_uppercase(
        self, document: Union[DocumentObject, _Cell]
    ):
        """Extract formatting and convert text to uppercase
        Args:
            document (Union[DocumentObject, _Cell]): Document or Cell object

        Returns:
            list: List of paragraphs with formatting information
        """
        paragraphs = []
        for para in document.paragraphs:
            if not para.text or para.text.isspace():
                continue
            para_info = {"paragraph": para.text, "runs": []}
            for run in para.runs:
                font = run.font
                para_info["runs"].append(
                    {
                        "content": run.text,
                        "font": font.name,
                        "font_size": font.size.pt if font.size else None,
                        "color": font.color.rgb if font.color else None,
                        "bold": font.bold,
                        "italic": font.italic,
                    }
                )
                run.text = run.text.upper()

            paragraphs.append(para_info)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(
                        self.extract_formatting_and_convert_uppercase(cell)
                    )

        return paragraphs

    def extract_convert_export(self, filename: str):
        """Main method to extract, convert and export the docx file"""
        img_dir = prepare_output_dir(filename, self.output_dir)

        # Extract images
        self.extract_images(filename, img_dir)

        document = Document(path.join(self.data_dir, filename))
        paragraphs = []

        # Extract formatting and convert to uppercase
        paragraphs.extend(self.extract_formatting_and_convert_uppercase(document))

        filename = filename.split(".")[0]
        # Export the extracted formatting data
        with open(path.join(self.output_dir, filename, f"{filename}.json"), "w") as f:
            json.dump(paragraphs, f, indent=2)

        # Export the converted docx file
        document.save(
            path.join(self.output_dir, filename, f"{filename}_uppercase.docx")
        )

        return paragraphs
